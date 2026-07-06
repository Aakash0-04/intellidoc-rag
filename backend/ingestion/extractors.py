"""PDF and DOCX extractors with rich metadata."""
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import List, Tuple
from datetime import datetime

import fitz
import pdfplumber
from docx import Document as DocxDocument


@dataclass
class ExtractedContent:
    """Unified output from any document loader."""
    text: str
    tables: List[str]
    metadata: dict
    text_pages: List[Tuple[int, str]]


class PDFExtractor:
    def extract(self, file_path: str | Path) -> ExtractedContent:
        file_path = Path(file_path)
        doc = fitz.open(file_path)
        
        text_pages = []
        full_text = []
        for page_num, page in enumerate(doc, start=1):
            text = page.get_text()
            if text.strip():
                text_pages.append((page_num, text.strip()))
                full_text.append(f"[Page {page_num}]\n{text.strip()}")
        
        meta = doc.metadata or {}
        doc.close()
        
        # Extract tables with page numbers
        tables = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, start=1):
                for table in (page.extract_tables() or []):
                    if table:
                        rows = [" | ".join(str(cell or "") for cell in row) for row in table]
                        tables.append((page_num, "\n".join(rows)))
        
        # Rich metadata
        return ExtractedContent(
            text="\n\n".join(full_text),
            tables=[t[1] for t in tables],
            metadata={
                "source": file_path.name,
                "file_type": "pdf",
                "title": meta.get("title") or file_path.stem,
                "author": meta.get("author"),
                "subject": meta.get("subject"),
                "keywords": meta.get("keywords"),
                "creator": meta.get("creator"),
                "producer": meta.get("producer"),
                "creation_date": meta.get("creationDate"),
                "modification_date": meta.get("modDate"),
                "total_pages": len(text_pages),
                "has_tables": len(tables) > 0,
                "file_size_kb": round(file_path.stat().st_size / 1024, 2),
                "extracted_at": datetime.utcnow().isoformat(),
            },
            text_pages=text_pages,
        )


class DOCXExtractor:
    def extract(self, file_path: str | Path) -> ExtractedContent:
        file_path = Path(file_path)
        doc = DocxDocument(file_path)
        
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables.append("\n".join(rows))
        
        core = doc.core_properties
        
        # Rich metadata
        return ExtractedContent(
            text="\n\n".join(paragraphs),
            tables=tables,
            metadata={
                "source": file_path.name,
                "file_type": "docx",
                "title": core.title or file_path.stem,
                "author": core.author,
                "subject": core.subject,
                "keywords": core.keywords,
                "created": core.created.isoformat() if core.created else None,
                "modified": core.modified.isoformat() if core.modified else None,
                "last_modified_by": core.last_modified_by,
                "total_paragraphs": len(paragraphs),
                "has_tables": len(tables) > 0,
                "file_size_kb": round(file_path.stat().st_size / 1024, 2),
                "extracted_at": datetime.utcnow().isoformat(),
            },
            text_pages=[(1, "\n\n".join(paragraphs))] if paragraphs else [],
        )